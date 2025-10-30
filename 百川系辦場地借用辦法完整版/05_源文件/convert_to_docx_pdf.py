#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to DOCX and PDF Converter
將 Markdown 文件轉換為 DOCX 和 PDF 格式
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
import os
import sys

def read_markdown_file(filepath):
    """讀取 Markdown 文件"""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def convert_to_docx(markdown_text, output_path):
    """將 Markdown 轉換為 DOCX"""
    doc = Document()

    # 設定文檔樣式
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)

    # 分行處理
    lines = markdown_text.split('\n')

    for line in lines:
        line = line.strip()

        if not line:
            continue

        # 處理標題
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)

        # 處理表格（簡化版，檢測表格標記）
        elif line.startswith('|'):
            # 表格處理（簡化）
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            table_data = [cells]
            doc.add_paragraph(' | '.join(cells))

        # 處理列表
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s*', '', line)
            doc.add_paragraph(text, style='List Number')

        # 處理代碼塊標記
        elif line.startswith('```'):
            continue  # 簡化處理，跳過代碼塊標記

        # 處理分隔線
        elif line.startswith('---') or line.startswith('***'):
            p = doc.add_paragraph()
            p.add_run('─' * 80)

        # 普通段落
        else:
            # 移除 Markdown 格式標記（簡化）
            text = line
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # 粗體
            text = re.sub(r'\*(.*?)\*', r'\1', text)      # 斜體
            text = re.sub(r'`(.*?)`', r'\1', text)        # 代碼
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # 連結

            if text:
                doc.add_paragraph(text)

    # 儲存文件
    doc.save(output_path)
    print(f"[OK] DOCX file created: {output_path}")

def convert_to_pdf(markdown_text, output_path):
    """將 Markdown 轉換為 PDF（簡化版）"""

    # 創建 PDF 文件
    pdf = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=18
    )

    # 準備故事內容
    story = []
    styles = getSampleStyleSheet()

    # 自定義樣式（處理中文）
    try:
        # 嘗試使用系統字體（Windows）
        font_path = 'C:\\Windows\\Fonts\\msyh.ttc'  # 微軟雅黑
        if os.path.exists(font_path):
            pdfmetrics.registerFont(TTFont('Chinese', font_path))

            # 創建中文樣式
            chinese_style = ParagraphStyle(
                'Chinese',
                parent=styles['Normal'],
                fontName='Chinese',
                fontSize=11,
                leading=16
            )

            chinese_heading1 = ParagraphStyle(
                'ChineseHeading1',
                parent=styles['Heading1'],
                fontName='Chinese',
                fontSize=18,
                leading=22,
                textColor=colors.HexColor('#1a1a1a')
            )

            chinese_heading2 = ParagraphStyle(
                'ChineseHeading2',
                parent=styles['Heading2'],
                fontName='Chinese',
                fontSize=14,
                leading=18,
                textColor=colors.HexColor('#333333')
            )
        else:
            chinese_style = styles['Normal']
            chinese_heading1 = styles['Heading1']
            chinese_heading2 = styles['Heading2']
    except Exception as e:
        print(f"[WARNING] Cannot load Chinese font, using default. Error: {e}")
        chinese_style = styles['Normal']
        chinese_heading1 = styles['Heading1']
        chinese_heading2 = styles['Heading2']

    # 分行處理
    lines = markdown_text.split('\n')

    for line in lines:
        line = line.strip()

        if not line:
            story.append(Spacer(1, 0.2*inch))
            continue

        # 處理標題
        if line.startswith('# '):
            text = line[2:]
            story.append(Paragraph(text, chinese_heading1))
            story.append(Spacer(1, 0.3*inch))
        elif line.startswith('## '):
            text = line[3:]
            story.append(Paragraph(text, chinese_heading2))
            story.append(Spacer(1, 0.2*inch))
        elif line.startswith('### ') or line.startswith('#### '):
            text = line.split(' ', 1)[1] if ' ' in line else ''
            story.append(Paragraph(f"<b>{text}</b>", chinese_style))
            story.append(Spacer(1, 0.1*inch))

        # 處理分隔線
        elif line.startswith('---') or line.startswith('***'):
            story.append(Spacer(1, 0.1*inch))
            story.append(Paragraph('_' * 80, styles['Normal']))
            story.append(Spacer(1, 0.1*inch))

        # 跳過代碼塊標記
        elif line.startswith('```'):
            continue

        # 跳過表格分隔線
        elif line.startswith('|') and '-' in line:
            continue

        # 處理表格行
        elif line.startswith('|'):
            continue  # 簡化處理，暫時跳過表格

        # 處理列表
        elif line.startswith('- ') or line.startswith('* '):
            text = '• ' + line[2:]
            story.append(Paragraph(text, chinese_style))
        elif re.match(r'^\d+\.', line):
            story.append(Paragraph(line, chinese_style))

        # 普通段落
        else:
            # 移除 Markdown 格式標記
            text = line
            text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)  # 粗體
            text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)      # 斜體
            text = re.sub(r'`(.*?)`', r'<font face="Courier">\1</font>', text)  # 代碼
            text = re.sub(r'\[(.*?)\]\((.*?)\)', r'<u>\1</u>', text)  # 連結

            if text:
                try:
                    story.append(Paragraph(text, chinese_style))
                except Exception as e:
                    # 如果段落處理失敗，嘗試純文本
                    print(f"[WARNING] Paragraph processing failed, using plain text. Error: {e}")
                    clean_text = re.sub(r'<.*?>', '', text)
                    story.append(Paragraph(clean_text, chinese_style))

    # 生成 PDF
    try:
        pdf.build(story)
        print(f"[OK] PDF file created: {output_path}")
    except Exception as e:
        print(f"[ERROR] PDF generation failed: {e}")
        print("Note: PDF may contain special characters or formats that cannot be processed")

def main():
    """主函數"""
    print("=" * 80)
    print("Markdown to DOCX & PDF Converter")
    print("=" * 80)
    print()

    # 要轉換的文件
    files_to_convert = [
        "國立陽明交通大學百川學士學位學程系辦場地借用辦法_彈性版.md",
        "國立陽明交通大學百川學士學位學程系辦場地借用辦法.md",
        "百川系辦場地借用_最佳實踐分析.md",
        "簽呈_訂定百川系辦場地借用辦法.md"
    ]

    base_dir = os.path.dirname(os.path.abspath(__file__))

    for filename in files_to_convert:
        filepath = os.path.join(base_dir, filename)

        if not os.path.exists(filepath):
            print(f"[SKIP] File not found: {filename}")
            continue

        print(f"\nProcessing: {filename}")
        print("-" * 80)

        # 讀取 Markdown
        try:
            markdown_text = read_markdown_file(filepath)
        except Exception as e:
            print(f"[ERROR] Failed to read file: {e}")
            continue

        # 生成輸出文件名
        base_name = os.path.splitext(filename)[0]
        docx_path = os.path.join(base_dir, f"{base_name}.docx")
        pdf_path = os.path.join(base_dir, f"{base_name}.pdf")

        # 轉換為 DOCX
        try:
            convert_to_docx(markdown_text, docx_path)
        except Exception as e:
            print(f"[ERROR] DOCX conversion failed: {e}")

        # 轉換為 PDF
        try:
            convert_to_pdf(markdown_text, pdf_path)
        except Exception as e:
            print(f"[ERROR] PDF conversion failed: {e}")

    print("\n" + "=" * 80)
    print("Conversion Complete!")
    print("=" * 80)

if __name__ == "__main__":
    main()
